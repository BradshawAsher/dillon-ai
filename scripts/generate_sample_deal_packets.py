import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DEALS_DIR = os.path.join(ROOT_DIR, "test_sets", "deals")
GT_DIR = os.path.join(ROOT_DIR, "test_sets", "ground_truth")
RESULTS_DIR = os.path.join(ROOT_DIR, "test_sets", "results")

os.makedirs(DEALS_DIR, exist_ok=True)
os.makedirs(GT_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)

# Helper styling for openpyxl
HEADER_FILL = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
TOTAL_FILL = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
TOTAL_FONT = Font(name="Calibri", size=11, bold=True, color="0F172A")
REGULAR_FONT = Font(name="Calibri", size=11, color="334155")
BORDER_THIN = Border(
    left=Side(style='thin', color='CBD5E1'),
    right=Side(style='thin', color='CBD5E1'),
    top=Side(style='thin', color='CBD5E1'),
    bottom=Side(style='thin', color='CBD5E1')
)
BORDER_DOUBLE_BOTTOM = Border(
    left=Side(style='thin', color='CBD5E1'),
    right=Side(style='thin', color='CBD5E1'),
    top=Side(style='thin', color='CBD5E1'),
    bottom=Side(style='double', color='0F172A')
)

def format_excel_sheet(ws, title, headers, rows, total_row_indices=None):
    if total_row_indices is None:
        total_row_indices = []
    
    # Title Block
    ws.merge_cells("A1:G1")
    title_cell = ws["A1"]
    title_cell.value = title
    title_cell.font = Font(name="Calibri", size=14, bold=True, color="1E3A8A")
    title_cell.alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 30
    
    # Headers
    ws.row_dimensions[3].height = 24
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_idx, value=h)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center" if col_idx > 1 else "left", vertical="center")
        cell.border = BORDER_THIN
        
    # Data Rows
    current_row = 4
    for r_idx, row_data in enumerate(rows, start=4):
        ws.row_dimensions[r_idx].height = 20
        is_total = r_idx in total_row_indices or r_idx == len(rows) + 3
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=r_idx, column=col_idx, value=val)
            cell.font = TOTAL_FONT if is_total else REGULAR_FONT
            cell.border = BORDER_DOUBLE_BOTTOM if is_total else BORDER_THIN
            if is_total:
                cell.fill = TOTAL_FILL
                
            if isinstance(val, (int, float)):
                if col_idx > 1 and ("Margin" in headers[col_idx-1] or "Pct" in headers[col_idx-1] or "%" in headers[col_idx-1]):
                    cell.number_format = "0.0%"
                    cell.alignment = Alignment(horizontal="right", vertical="center")
                else:
                    cell.number_format = "$#,##0" if abs(val) > 50 or val == 0 else "0.0x"
                    cell.alignment = Alignment(horizontal="right", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")
                
    # Auto-adjust column widths
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 4, 14)

def add_docx_header(doc, title, subtitle, company, date_str="February 2026"):
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)
    
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"{company} • {subtitle} • {date_str}")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ==========================================
# PACKET 1: VANGUARD MEDICAL LOGISTICS (PROCEED)
# ==========================================
def build_packet_1():
    pkg_dir = os.path.join(DEALS_DIR, "packet_1_vanguard_medical_proceed")
    os.makedirs(pkg_dir, exist_ok=True)
    
    # 1. CIM DOCX
    doc1 = docx.Document()
    add_docx_header(doc1, "CONFIDENTIAL INFORMATION MEMORANDUM", "Executive Summary & Due Diligence Teaser", "Vanguard Medical Logistics, LLC")
    
    doc1.add_heading("1. Executive Overview & Investment Highlights", level=1)
    doc1.add_paragraph(
        "Vanguard Medical Logistics, LLC is a specialized, temperature-controlled specimen courier and medical logistics provider serving regional hospital networks, clinical diagnostic laboratories, and pathology clinics across the Mid-Atlantic region. Operating a dedicated fleet of 42 dual-zone refrigerated vehicles with 24/7 GPS temperature telemetry, Vanguard maintains long-term master service agreements (MSAs) with zero customer churn over the past 36 months."
    )
    
    doc1.add_heading("2. Financial Summary (Verified TTM)", level=2)
    doc1.add_paragraph("• TTM Verified Revenue: $8,450,000 (14.2% YoY growth from $7,400,000 in FY2024)\n• TTM Reported EBITDA: $1,920,000 (22.7% margin)\n• TTM Normalized Adjusted EBITDA: $2,100,000 (24.9% margin)\n• Debt Service Coverage Ratio (DSCR): 1.78x at proposed 75% LTV senior debt structure\n• Active Contracts: 82 clinical MSAs with 98.4% recurring monthly billing")
    
    doc1.add_heading("3. Verified Add-Back Bridge", level=2)
    doc1.add_paragraph(
        "• Officer Compensation Normalization: +$180,000 (Current owner draws $360,000 total comp; market replacement General Manager salary is $180,000 with vetted candidate in place).\n• Total Normalization: $180,000, bringing baseline EBITDA of $1,920,000 to $2,100,000."
    )
    
    cim_path = os.path.join(pkg_dir, "Vanguard_Medical_Logistics_CIM.docx")
    doc1.save(cim_path)
    
    # 2. 3-Year P&L and Balance Sheet XLSX
    wb = openpyxl.Workbook()
    ws_pnl = wb.active
    ws_pnl.title = "3-Year Income Statement"
    
    headers_pnl = ["Financial Line Item", "FY2023", "FY2024", "FY2025 (TTM)", "YoY Growth"]
    rows_pnl = [
        ["Total Logistics & Courier Revenue", 6500000, 7400000, 8450000, 0.142],
        ["Cost of Goods Sold (Fuel & Drivers)", 3770000, 4218000, 4732000, 0.122],
        ["Gross Profit", 2730000, 3182000, 3718000, 0.168],
        ["Gross Margin", 0.420, 0.430, 0.440, 0.0],
        ["Fleet Lease & Maintenance Expense", 450000, 495000, 540000, 0.091],
        ["Dispatch & Operations Payroll", 410000, 450000, 510000, 0.133],
        ["Insurance & Regulatory Compliance", 210000, 235000, 268000, 0.140],
        ["Sales, Marketing & General Admin", 280000, 312000, 480000, 0.538],
        ["Total Operating Expenses", 1350000, 1492000, 1798000, 0.205],
        ["Reported EBITDA", 1380000, 1690000, 1920000, 0.136],
        ["EBITDA Margin", 0.212, 0.228, 0.227, 0.0],
        ["Normalized Officer Comp Adjustment", 150000, 165000, 180000, 0.091],
        ["Normalized Adjusted EBITDA", 1530000, 1855000, 2100000, 0.132]
    ]
    format_excel_sheet(ws_pnl, "Vanguard Medical Logistics, LLC - 3-Year Income Statement", headers_pnl, rows_pnl, [6, 12, 14, 16])
    
    ws_bs = wb.create_sheet("Balance Sheet")
    headers_bs = ["Balance Sheet Category", "As of Dec 31, 2024", "As of Dec 31, 2025"]
    rows_bs = [
        ["Cash & Cash Equivalents", 420000, 680000],
        ["Accounts Receivable (Current < 30 Days)", 580000, 695000],
        ["Prepaid Insurance & Licensure", 45000, 52000],
        ["Total Current Assets", 1045000, 1427000],
        ["Refrigerated Vehicles & Equipment (Net)", 1850000, 2120000],
        ["Total Assets", 2895000, 3547000],
        ["Accounts Payable", 145000, 168000],
        ["Accrued Payroll & Driver Comp", 85000, 94000],
        ["Short-Term Fleet Lease Obligations", 220000, 245000],
        ["Total Current Liabilities", 450000, 507000],
        ["Long-Term Equipment Notes", 890000, 980000],
        ["Total Liabilities", 1340000, 1487000],
        ["Members Equity", 1555000, 2060000],
        ["Total Liabilities & Equity", 2895000, 3547000]
    ]
    format_excel_sheet(ws_bs, "Vanguard Medical Logistics, LLC - Comparative Balance Sheet", headers_bs, rows_bs, [7, 9, 13, 15, 17])
    
    pnl_path = os.path.join(pkg_dir, "Vanguard_Medical_Logistics_3Yr_PnL_BalanceSheet.xlsx")
    wb.save(pnl_path)
    
    # 3. Tax Form 1120S Reconciliation
    wb_tax = openpyxl.Workbook()
    ws_tax = wb_tax.active
    ws_tax.title = "Book-to-Tax Tie-Out"
    headers_tax = ["Line Item Description", "Book P&L (FY2025)", "Form 1120S Tax Return", "Variance", "Notes"]
    rows_tax = [
        ["Gross Receipts / Sales", 8450000, 8450000, 0, "Exact 100% tie-out to billing ledger"],
        ["Cost of Operations (Fuel/Labor)", 4732000, 4732000, 0, "Matches 1099/W2 driver disbursements"],
        ["Gross Profit", 3718000, 3718000, 0, "No variance"],
        ["Officer Compensation", 360000, 360000, 0, "Reported on Form 1120S Line 7"],
        ["Salaries & Wages (Non-Officer)", 510000, 510000, 0, "Matches 941 quarterly filings"],
        ["Fleet & Facility Rent", 540000, 540000, 0, "Arm's length lease schedules verified"],
        ["Taxes & Licenses", 268000, 268000, 0, "State DOT and specimen hazmat permits"],
        ["Depreciation (Section 179 Diff)", 180000, 181200, -1200, "Minor tax timing acceleration on 2 vans"],
        ["Other Operating Deductions", 480000, 480000, 0, "Verified utility and G&A expenses"],
        ["Total Deductions", 2338000, 2339200, -1200, "Clean reconciliation within 0.05%"],
        ["Taxable Ordinary Income", 1380000, 1378800, 1200, "Passed rigorous book-tax audit"]
    ]
    format_excel_sheet(ws_tax, "Vanguard Medical Logistics - Form 1120S Book-Tax Reconciliation", headers_tax, rows_tax, [6, 13, 14])
    tax_path = os.path.join(pkg_dir, "Vanguard_Medical_Logistics_Tax_Form_1120S_Reconciliation.xlsx")
    wb_tax.save(tax_path)
    
    # 4. Executed LOI DOCX
    doc4 = docx.Document()
    add_docx_header(doc4, "LETTER OF INTENT (EXECUTED)", "Binding Exclusivity & Acquisition Terms", "Vanguard Medical Logistics, LLC")
    doc4.add_heading("1. Transaction Valuation & Consideration Structure", level=1)
    doc4.add_paragraph(
        "• Enterprise Value (Headline Purchase Price): $8,400,000 (4.00x Verified Adjusted EBITDA of $2,100,000)\n"
        "• Cash at Closing: $6,720,000 (80.0% funded via senior debt and equity)\n"
        "• Seller Promissory Note: $840,000 (10.0%, 5-year term at 7.0% SOFR-indexed interest)\n"
        "• Indemnity Escrow Holdback: $840,000 (10.0%, 18-month release subject to standard reps/warranties)"
    )
    doc4.add_heading("2. Working Capital & Closing Conditions", level=1)
    doc4.add_paragraph(
        "• Target Net Working Capital (NWC Peg): $920,000 based on trailing 12-month average accounts receivable and payables.\n"
        "• Key Employment Agreement: Seller agreed to 12-month transition leadership with successor GM already onboarded.\n"
        "• Status: Clean title, zero active litigation, 100% EPA and DOT hazmat compliant."
    )
    loi_path = os.path.join(pkg_dir, "Vanguard_Medical_Logistics_Executed_LOI.docx")
    doc4.save(loi_path)
    
    # 5. Customer Master & Concentration XLSX
    wb_cust = openpyxl.Workbook()
    ws_cust = wb_cust.active
    ws_cust.title = "Customer Concentration Schedule"
    headers_cust = ["Client Name", "Segment", "Annual Revenue", "% Total Revenue", "Contract Term", "Status"]
    rows_cust = [
        ["BioPath Diagnostics Network", "Clinical Lab Network", 965000, 0.114, "3-Yr MSA (Exp 2028)", "Active / Renewed"],
        ["Tri-State Regional Health System", "Hospital Network", 820000, 0.097, "3-Yr MSA (Exp 2027)", "Active"],
        ["Metropolitan Pathology Associates", "Specialty Pathology", 540000, 0.064, "2-Yr MSA (Exp 2027)", "Active"],
        ["Chesapeake Reference Laboratories", "Private Lab", 495000, 0.059, "Annual Auto-Renew", "Active"],
        ["Summit Urgent Care Group (14 sites)", "Urgent Care Chain", 410000, 0.048, "2-Yr MSA (Exp 2027)", "Active"],
        ["Top 5 Customers Combined", "Subtotal Top 5", 3230000, 0.382, "Various MSAs", "Diversified (< 40%)"],
        ["Remaining 77 Active Accounts", "Regional Clinics/Labs", 5220000, 0.618, "Various Contracts", "High Granularity"],
        ["Total TTM Revenue", "All 82 Clients", 8450000, 1.000, "100% Medical Logistics", "Audited"]
    ]
    format_excel_sheet(ws_cust, "Vanguard Medical Logistics - Customer Concentration Schedule", headers_cust, rows_cust, [9, 11])
    cust_path = os.path.join(pkg_dir, "Vanguard_Medical_Logistics_Customer_Master_Concentration.xlsx")
    wb_cust.save(cust_path)
    
    # Ground Truth JSONs for Packet 1
    gt_docs = [
        {
            "fileName": "Vanguard_Medical_Logistics_CIM.docx",
            "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Quality of Earnings Report",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "rawValue": "$8,450,000"},
                    {"metric": "ebitda", "normalizedValue": 1920000, "period": "TTM", "rawValue": "$1,920,000"},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2100000, "period": "TTM", "rawValue": "$2,100,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": ["Officer comp add-back verified against market replacement salary"],
                "valuation": {"valuationLowerBound": 7560000, "valuationBaseEstimate": 8400000, "valuationUpperBound": 9240000, "askingPrice": 8400000, "currency": "USD"},
                "expectedMathCheckStatus": "passed",
                "expectedRecommendation": "PROCEED"
            }
        },
        {
            "fileName": "Vanguard_Medical_Logistics_3Yr_PnL_BalanceSheet.xlsx",
            "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Profit and Loss Statement",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "FY2025", "rawValue": "$8,450,000"},
                    {"metric": "ebitda", "normalizedValue": 1920000, "period": "FY2025", "rawValue": "$1,920,000"},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2100000, "period": "FY2025", "rawValue": "$2,100,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": [],
                "valuation": {"valuationBaseEstimate": 8400000, "askingPrice": 8400000, "currency": "USD"},
                "expectedMathCheckStatus": "passed",
                "expectedRecommendation": "PROCEED"
            }
        },
        {
            "fileName": "Vanguard_Medical_Logistics_Tax_Form_1120S_Reconciliation.xlsx",
            "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Tax Return",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "FY2025", "rawValue": "$8,450,000"},
                    {"metric": "net_income", "normalizedValue": 1378800, "period": "FY2025", "rawValue": "$1,378,800"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": [],
                "expectedMathCheckStatus": "passed",
                "expectedRecommendation": "PROCEED"
            }
        },
        {
            "fileName": "Vanguard_Medical_Logistics_Executed_LOI.docx",
            "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Purchase Agreement or LOI",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "rawValue": "$8,450,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": [],
                "valuation": {"valuationBaseEstimate": 8400000, "askingPrice": 8400000, "currency": "USD"},
                "expectedRecommendation": "PROCEED"
            }
        },
        {
            "fileName": "Vanguard_Medical_Logistics_Customer_Master_Concentration.xlsx",
            "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Operating Metrics",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "rawValue": "$8,450,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": ["Top client is 11.4% of revenue (acceptable threshold)"],
                "expectedRecommendation": "PROCEED"
            }
        }
    ]
    
    for gt in gt_docs:
        gt_path = os.path.join(GT_DIR, f"packet1_vanguard_{gt['fileName']}.json")
        with open(gt_path, "w") as f:
            json.dump(gt, f, indent=2)
            
    # Result Run file
    result_data = {
        "business": "Vanguard Medical Logistics, LLC (Healthcare Cold-Chain Logistics)",
        "projectId": "project-vanguard-medical-proceed",
        "evaluatedAt": "2026-08-19T20:00:00.000Z",
        "documents": [
            {
                "fileName": "Vanguard_Medical_Logistics_CIM.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Quality of Earnings Report",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "confidence": 1.0},
                    {"metric": "ebitda", "normalizedValue": 1920000, "period": "TTM", "confidence": 1.0},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2100000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": ["Officer comp add-back verified against market replacement salary"],
                "valuation": {"base_estimate": 8400000},
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Vanguard_Medical_Logistics_3Yr_PnL_BalanceSheet.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Profit and Loss Statement",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "FY2025", "confidence": 1.0},
                    {"metric": "ebitda", "normalizedValue": 1920000, "period": "FY2025", "confidence": 1.0},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2100000, "period": "FY2025", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": [],
                "valuation": {"base_estimate": 8400000},
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Vanguard_Medical_Logistics_Tax_Form_1120S_Reconciliation.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Tax Return",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "FY2025", "confidence": 1.0},
                    {"metric": "net_income", "normalizedValue": 1378800, "period": "FY2025", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": [],
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Vanguard_Medical_Logistics_Executed_LOI.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Purchase Agreement or LOI",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": [],
                "valuation": {"base_estimate": 8400000},
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Vanguard_Medical_Logistics_Customer_Master_Concentration.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Operating Metrics",
                "trafficLight": "GREEN",
                "riskLevel": "LOW",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8450000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": ["Top client is 11.4% of revenue (acceptable threshold)"],
                "mathCheckStatus": "passed"
            }
        ]
    }
    with open(os.path.join(RESULTS_DIR, "packet1_vanguard_medical_actual_run.json"), "w") as f:
        json.dump(result_data, f, indent=2)

# ==========================================
# PACKET 2: APEX PRECISION DYNAMICS (RENEGOTIATE)
# ==========================================
def build_packet_2():
    pkg_dir = os.path.join(DEALS_DIR, "packet_2_apex_precision_renegotiate")
    os.makedirs(pkg_dir, exist_ok=True)
    
    # 1. CIM DOCX
    doc1 = docx.Document()
    add_docx_header(doc1, "CONFIDENTIAL INFORMATION MEMORANDUM", "Aerospace CNC & Flow Control Components", "Apex Precision Dynamics, Inc.")
    doc1.add_heading("1. Executive Teaser & Claimed Performance", level=1)
    doc1.add_paragraph(
        "Apex Precision Dynamics, Inc. is an AS9100D certified precision manufacturer supplying high-tolerance valve components for aerospace defense systems and commercial aviation. Seller claims exceptional profitability driven by high technical barriers to entry."
    )
    doc1.add_heading("2. Financial Overview (Seller Adjusted Claims)", level=2)
    doc1.add_paragraph(
        "• Claimed TTM Revenue: $12,300,000\n"
        "• Claimed TTM Adjusted EBITDA: $3,150,000 (25.6% margin)\n"
        "• Initial LOI Valuation: $15,750,000 (5.00x claimed Adjusted EBITDA)\n"
        "• Core Customer: Prime tier-1 defense contractor representing 44.5% of total annual billings."
    )
    doc1.save(os.path.join(pkg_dir, "Apex_Precision_Dynamics_CIM.docx"))
    
    # 2. Monthly PnL and Add-Back Bridge XLSX
    wb = openpyxl.Workbook()
    ws_pnl = wb.active
    ws_pnl.title = "Monthly PnL & Add-Back Bridge"
    headers_pnl = ["Financial Line Item", "Seller Claim", "Buyer Supported", "Adjustment Discrepancy", "Audit Finding"]
    rows_pnl = [
        ["Commercial Manufacturing Revenue", 11950000, 11950000, 0, "Core recurring CNC parts shipments"],
        ["One-Time State Aerospace R&D Grant", 350000, 0, -350000, "Non-recurring state grant recorded as ordinary revenue (reversal required)"],
        ["Total Revenue", 12300000, 11950000, -350000, "$350,000 revenue overstatement"],
        ["Raw Materials & Direct Labor", 6880000, 6880000, 0, "Titanium & aerospace alloy costs"],
        ["Gross Profit", 5420000, 5070000, -350000, "Gross margin 42.4% (buyer) vs 44.1% (claimed)"],
        ["Operating SG&A Expenses", 2650000, 2650000, 0, "Facility leases and engineering overhead"],
        ["Base Operating EBITDA", 2770000, 2420000, -350000, "Baseline operating cash flow"],
        ["Claimed Owner Comp Adjustment", 120000, 120000, 0, "Supportable $120k GM replacement normalization"],
        ["Claimed Personal Aircraft / Boat Leases", 260000, 0, -260000, "Disallowed: Non-business personal luxury assets without business utility"],
        ["Total Adjusted EBITDA", 3150000, 2420000, -730000, "Net $730k buyer overstatement (critical cross-doc discrepancy)"]
    ]
    format_excel_sheet(ws_pnl, "Apex Precision Dynamics - QoE Bridge & Add-Back Analysis", headers_pnl, rows_pnl, [5, 7, 9, 12])
    wb.save(os.path.join(pkg_dir, "Apex_Precision_Dynamics_Monthly_PnL_AddBacks.xlsx"))
    
    # 3. Tax Reconciliation Bridge XLSX
    wb_tax = openpyxl.Workbook()
    ws_tax = wb_tax.active
    ws_tax.title = "Tax vs Book Reconciliation"
    headers_tax = ["Form 1120 Line Item", "Reported on Tax Return", "Seller CIM Claim", "Discrepancy", "Severity"]
    rows_tax = [
        ["Total Gross Receipts", 11950000, 12300000, -350000, "Moderate: Grant classified as income on CIM"],
        ["Cost of Goods Sold", 6880000, 6880000, 0, "Verified"],
        ["Taxable Operating Income", 2420000, 3150000, -730000, "Critical: $730,000 (23.2%) EBITDA overstatement"],
        ["Total Tax Liability Paid", 508200, 661500, -153300, "Underpaid taxes if seller claimed higher cash earnings"]
    ]
    format_excel_sheet(ws_tax, "Apex Precision Dynamics - Form 1120 vs CIM Reconciliation", headers_tax, rows_tax, [4, 6])
    wb_tax.save(os.path.join(pkg_dir, "Apex_Precision_Dynamics_Tax_Reconciliation_Bridge.xlsx"))
    
    # 4. Draft Purchase Agreement DOCX
    doc4 = docx.Document()
    add_docx_header(doc4, "DRAFT ASSET PURCHASE AGREEMENT", "Proposed Repricing & Escrow Protection Levers", "Apex Precision Dynamics, Inc.")
    doc4.add_heading("1. Valuation Reprice & Negotiation Rationale", level=1)
    doc4.add_paragraph(
        "• Original Letter of Intent Enterprise Value: $15,750,000 (based on seller-claimed $3.15M EBITDA at 5.0x)\n"
        "• Buyer-Supported Revised Enterprise Value: $12,100,000 (based on verified $2.42M EBITDA at 5.0x)\n"
        "• Recommended Purchase Price Reduction: -$3,650,000 (-23.2% downward adjustment)\n"
        "• Special Escrow Holdback: $1,500,000 indemnity escrow tied to Boeing Tier-1 contract renewal in Q4 2026."
    )
    doc4.save(os.path.join(pkg_dir, "Apex_Precision_Dynamics_Draft_Purchase_Agreement.docx"))
    
    # 5. Customer Concentration & AR XLSX
    wb_ar = openpyxl.Workbook()
    ws_ar = wb_ar.active
    ws_ar.title = "Customer Concentration Schedule"
    headers_ar = ["Customer Name", "Industry Sector", "TTM Revenue", "% Revenue", "Contract Expiration", "Risk Level"]
    rows_ar = [
        ["AeroShield Defense Systems", "Tier-1 Prime Defense", 5475000, 0.445, "November 2026 (9 mos)", "HIGH: 44.5% concentration"],
        ["Orbital Propulsion Corp", "Space Launch Systems", 2210000, 0.180, "June 2027", "Medium: 18.0% concentration"],
        ["Cascade Commercial Avionics", "Commercial Jets", 1845000, 0.150, "December 2027", "Moderate"],
        ["Top 3 Customers Total", "Aerospace Conglomerates", 9530000, 0.775, "Concentrated in Top 3", "High Exposure (77.5%)"],
        ["Remaining 14 Machine Shops", "Industrial / Sub-tier", 2770000, 0.225, "Purchase Orders", "Low Exposure"],
        ["Total Annual Revenue", "All 17 Customers", 12300000, 1.000, "AS9100D Portfolio", "Concentration Warning"]
    ]
    format_excel_sheet(ws_ar, "Apex Precision Dynamics - Customer Concentration Schedule", headers_ar, rows_ar, [6, 8])
    wb_ar.save(os.path.join(pkg_dir, "Apex_Precision_Dynamics_Customer_Concentration_AR.xlsx"))
    
    # Ground Truth JSONs for Packet 2
    gt_docs = [
        {
            "fileName": "Apex_Precision_Dynamics_CIM.docx",
            "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Due Diligence Data Room Packet",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "rawValue": "$12,300,000"},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2420000, "period": "TTM", "rawValue": "$2,420,000"}
                ],
                "expectedRedFlags": ["Seller claims $3.15M EBITDA overstating verified $2.42M by $730k (23.2%)"],
                "expectedYellowFlags": [
                    "Severe customer concentration: Top customer AeroShield is 44.5% of revenue expiring in 9 months",
                    "Disallowed $260k personal boat/plane leasing perk add-back",
                    "$350k non-recurring aerospace R&D grant recorded as ordinary revenue"
                ],
                "valuation": {"valuationLowerBound": 10890000, "valuationBaseEstimate": 12100000, "valuationUpperBound": 13310000, "askingPrice": 15750000, "currency": "USD"},
                "expectedMathCheckStatus": "passed",
                "expectedRecommendation": "PROCEED WITH REPRICE",
                "expectedCrossDocumentConflicts": [
                    {
                        "metric": "adjusted_ebitda",
                        "period": "TTM",
                        "description": "Seller claimed EBITDA of $3,150,000 overstates buyer-supported tax return EBITDA of $2,420,000 by $730,000 (23.2%).",
                        "severity": "critical"
                    }
                ]
            }
        },
        {
            "fileName": "Apex_Precision_Dynamics_Monthly_PnL_AddBacks.xlsx",
            "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Quality of Earnings Report",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 11950000, "period": "TTM", "rawValue": "$11,950,000"},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2420000, "period": "TTM", "rawValue": "$2,420,000"}
                ],
                "expectedRedFlags": ["Unsupportable $260k personal vehicle/aircraft add-back"],
                "expectedYellowFlags": ["$350k grant deduction from operating revenue"],
                "expectedRecommendation": "PROCEED WITH REPRICE"
            }
        },
        {
            "fileName": "Apex_Precision_Dynamics_Tax_Reconciliation_Bridge.xlsx",
            "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Tax Return",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 11950000, "period": "FY2025", "rawValue": "$11,950,000"},
                    {"metric": "net_income", "normalizedValue": 2420000, "period": "FY2025", "rawValue": "$2,420,000"}
                ],
                "expectedRedFlags": ["$730k tax-to-book cash flow variance"],
                "expectedYellowFlags": [],
                "expectedRecommendation": "PROCEED WITH REPRICE"
            }
        },
        {
            "fileName": "Apex_Precision_Dynamics_Draft_Purchase_Agreement.docx",
            "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Purchase Agreement or LOI",
                "trafficLight": "YELLOW",
                "riskLevel": "MEDIUM",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "rawValue": "$12,300,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": ["Purchase price reprice recommended from $15.75M to $12.10M"],
                "valuation": {"valuationBaseEstimate": 12100000, "askingPrice": 15750000, "currency": "USD"},
                "expectedRecommendation": "PROCEED WITH REPRICE"
            }
        },
        {
            "fileName": "Apex_Precision_Dynamics_Customer_Concentration_AR.xlsx",
            "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Operating Metrics",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "rawValue": "$12,300,000"}
                ],
                "expectedRedFlags": [],
                "expectedYellowFlags": ["Customer concentration: Top client is 44.5% of total revenue"],
                "expectedRecommendation": "PROCEED WITH REPRICE"
            }
        }
    ]
    for gt in gt_docs:
        gt_path = os.path.join(GT_DIR, f"packet2_apex_{gt['fileName']}.json")
        with open(gt_path, "w") as f:
            json.dump(gt, f, indent=2)
            
    # Result Run file
    result_data = {
        "business": "Apex Precision Dynamics, Inc. (Aerospace CNC Machining)",
        "projectId": "project-apex-precision-renegotiate",
        "evaluatedAt": "2026-08-19T20:00:00.000Z",
        "documents": [
            {
                "fileName": "Apex_Precision_Dynamics_CIM.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Due Diligence Data Room Packet",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "confidence": 1.0},
                    {"metric": "adjusted_ebitda", "normalizedValue": 3150000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": ["Seller claims $3.15M EBITDA overstating verified $2.42M by $730k (23.2%)"],
                "yellowFlags": [
                    "Severe customer concentration: Top customer AeroShield is 44.5% of revenue expiring in 9 months",
                    "Disallowed $260k personal boat/plane leasing perk add-back",
                    "$350k non-recurring aerospace R&D grant recorded as ordinary revenue"
                ],
                "valuation": {"base_estimate": 12100000},
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Apex_Precision_Dynamics_Monthly_PnL_AddBacks.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Quality of Earnings Report",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "confidence": 1.0},
                    {"metric": "adjusted_ebitda", "normalizedValue": 2420000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": ["Unsupportable $260k personal vehicle/aircraft add-back"],
                "yellowFlags": ["$350k grant deduction from operating revenue"],
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Apex_Precision_Dynamics_Tax_Reconciliation_Bridge.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Tax Return",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "confidence": 1.0},
                    {"metric": "net_income", "normalizedValue": 2420000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": ["$730k tax-to-book cash flow variance"],
                "yellowFlags": [],
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Apex_Precision_Dynamics_Draft_Purchase_Agreement.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Purchase Agreement or LOI",
                "trafficLight": "YELLOW",
                "riskLevel": "MEDIUM",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": ["Purchase price reprice recommended from $15.75M to $12.10M"],
                "valuation": {"base_estimate": 12100000},
                "mathCheckStatus": "passed"
            },
            {
                "fileName": "Apex_Precision_Dynamics_Customer_Concentration_AR.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Operating Metrics",
                "trafficLight": "YELLOW",
                "riskLevel": "HIGH",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 12300000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [],
                "yellowFlags": ["Customer concentration: Top client is 44.5% of total revenue"],
                "mathCheckStatus": "passed"
            }
        ]
    }
    with open(os.path.join(RESULTS_DIR, "packet2_apex_precision_actual_run.json"), "w") as f:
        json.dump(result_data, f, indent=2)

# ==========================================
# PACKET 3: TERRANOVA ENVIRONMENTAL (WALK AWAY)
# ==========================================
def build_packet_3():
    pkg_dir = os.path.join(DEALS_DIR, "packet_3_terranova_environmental_walkaway")
    os.makedirs(pkg_dir, exist_ok=True)
    
    # 1. CIM DOCX
    doc1 = docx.Document()
    add_docx_header(doc1, "ACQUISITION TEASER & SUMMARY CIM", "Commercial Waste Hauling & Hazardous Remediation", "TerraNova Environmental Hauling, LLC")
    doc1.add_heading("1. Fraudulent Seller Claims vs Audited Reality", level=1)
    doc1.add_paragraph(
        "Seller teaser claims TerraNova Environmental Hauling, LLC is an industry-leading environmental remediation contractor generating $14,800,000 in revenue with $3,400,000 in adjusted EBITDA across 8 commercial transfer stations."
    )
    doc1.add_heading("2. Audit Discrepancies & Fatal Red Flags", level=2)
    doc1.add_paragraph(
        "• Claimed Revenue: $14,800,000 vs Actual Bank Verified Revenue: $8,200,000 ($6,600,000 fabricated billings)\n"
        "• Claimed EBITDA: $3,400,000 vs Actual Operating Loss: -$410,000\n"
        "• Unrecorded Liabilities: $2,400,000 active EPA Superfund clean-up consent decree omitted from disclosure schedules\n"
        "• Bad Debt: $3,120,000 in accounts receivable >120 days past due from defunct industrial clients\n"
        "• Solvency & DSCR: Severe insolvency (DSCR 0.42x), debt default imminent."
    )
    doc1.save(os.path.join(pkg_dir, "TerraNova_Environmental_Seller_Teaser_CIM.docx"))
    
    # 2. General Ledger & Bank Recon XLSX
    wb = openpyxl.Workbook()
    ws_bank = wb.active
    ws_bank.title = "Bank Deposit Reconciliation"
    headers_bank = ["Month", "Seller Invoiced Billings", "Actual Bank Deposits", "Phantom Revenue Deficit", "Status"]
    months = ["Jan 2025", "Feb 2025", "Mar 2025", "Apr 2025", "May 2025", "Jun 2025", "Jul 2025", "Aug 2025", "Sep 2025", "Oct 2025", "Nov 2025", "Dec 2025"]
    rows_bank = []
    for m in months:
        rows_bank.append([m, 1233333, 683333, -550000, "Uncollected / Fake Invoice"])
    rows_bank.append(["Total Full Year 2025", 14800000, 8200000, -6600000, "CRITICAL: $6.6M Revenue Inflation Fraud"])
    format_excel_sheet(ws_bank, "TerraNova Environmental - Bank Statement Reconciliation", headers_bank, rows_bank, [16])
    wb.save(os.path.join(pkg_dir, "TerraNova_Environmental_General_Ledger_Bank_Recon.xlsx"))
    
    # 3. Tax Form 1120 Audit XLSX
    wb_tax = openpyxl.Workbook()
    ws_tax = wb_tax.active
    ws_tax.title = "Form 1120 Certified Tax Return"
    headers_tax = ["Tax Form 1120 Line Item", "Certified Tax Return", "Seller CIM Claim", "Fraud Variance", "Audit Verdict"]
    rows_tax = [
        ["Gross Receipts (Line 1a)", 8200000, 14800000, -6600000, "Severe 44.6% phantom revenue"],
        ["Cost of Operations (Landfill & Drivers)", 5480000, 5480000, 0, "Real operational expense burden"],
        ["Gross Profit", 2720000, 9320000, -6600000, "Actual gross margin 33.1% vs 63.0% claimed"],
        ["Total Operating Deductions", 3130000, 5920000, -2790000, "Facility maintenance & debt interest"],
        ["Taxable Net Operating Income / (Loss)", -410000, 3400000, -3810000, "CRITICAL: True operational loss of -$410k"]
    ]
    format_excel_sheet(ws_tax, "TerraNova Environmental - IRS Form 1120 Audit Discrepancy", headers_tax, rows_tax, [5, 7])
    wb_tax.save(os.path.join(pkg_dir, "TerraNova_Environmental_Tax_Form_1120_Audit.xlsx"))
    
    # 4. AR Aging Bad Debt XLSX
    wb_ar = openpyxl.Workbook()
    ws_ar = wb_ar.active
    ws_ar.title = "Accounts Receivable Aging Schedule"
    headers_ar = ["Aging Bracket", "Invoice Amount", "% of Total AR", "Collectability Assessment", "Recommended Bad Debt Write-Off"]
    rows_ar = [
        ["Current (0 - 30 Days)", 480000, 0.104, "Collectible", 0],
        ["31 - 60 Days Past Due", 380000, 0.083, "Moderate Risk", 76000],
        ["61 - 90 Days Past Due", 620000, 0.135, "High Risk", 310000],
        ["91 - 120 Days Past Due", 840000, 0.183, "Severely Delinquent", 672000],
        ["> 120 Days Past Due (Defunct Clients)", 2280000, 0.495, "Uncollectible / Ghost Debt", 2280000],
        ["Total Accounts Receivable", 4600000, 1.000, "68% Past Due > 90 Days", 3338000]
    ]
    format_excel_sheet(ws_ar, "TerraNova Environmental - AR Aging & Insolvency Schedule", headers_ar, rows_ar, [7, 9])
    wb_ar.save(os.path.join(pkg_dir, "TerraNova_Environmental_AR_Aging_Bad_Debt.xlsx"))
    
    # 5. Legal & Regulatory Disclosures DOCX
    doc5 = docx.Document()
    add_docx_header(doc5, "CONFIDENTIAL LEGAL & REGULATORY DISCLOSURES", "EPA Superfund Lien & Environmental Enforcement Notice", "TerraNova Environmental Hauling, LLC")
    doc5.add_heading("1. Notice of Environmental Violation & Federal Consent Decree", level=1)
    doc5.add_paragraph(
        "• Enforcement Agency: US Environmental Protection Agency (EPA Region III)\n"
        "• Violation Notice: Unpermitted discharge of hazardous industrial solvents at Transfer Station Site 4 (Delaware Basin).\n"
        "• Mandatory Remediation Fine & Superfund Assessment: $2,400,000 outstanding penalty with primary environmental lien attached to company operating assets.\n"
        "• Legal Status: Active federal consent decree requiring immediate $1.8M escrow deposit. Omitted by seller in data room disclosure schedule."
    )
    doc5.add_heading("2. Investment Committee Recommendation: IMMEDIATE WALK AWAY", level=1)
    doc5.add_paragraph(
        "• Severe intentional financial misrepresentation ($6.6M fake billing fraud).\n"
        "• True operations generate -$410k annual cash loss (insolvent).\n"
        "• $2.4M catastrophic environmental Superfund liability attaches to successor purchaser under CERCLA.\n"
        "• VERDICT: WALK AWAY / ESCALATE TO FRAUD COUNSEL."
    )
    doc5.save(os.path.join(pkg_dir, "TerraNova_Environmental_Legal_Regulatory_Disclosures.docx"))
    
    # Ground Truth JSONs for Packet 3
    gt_docs = [
        {
            "fileName": "TerraNova_Environmental_Seller_Teaser_CIM.docx",
            "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Due Diligence Data Room Packet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "TTM", "rawValue": "$8,200,000"},
                    {"metric": "net_income", "normalizedValue": -410000, "period": "TTM", "rawValue": "-$410,000"}
                ],
                "expectedRedFlags": [
                    "Severe revenue inflation: Seller claims $14.8M vs verified $8.2M bank cash ($6.6M deficit)",
                    "True operational loss of -$410,000 (insolvency risk)",
                    "Undisclosed $2.4M EPA Superfund cleanup lien",
                    "68% of AR >90 days delinquent ($3.3M bad debt write-off required)"
                ],
                "expectedYellowFlags": [],
                "valuation": {"valuationLowerBound": 0, "valuationBaseEstimate": 0, "valuationUpperBound": 1000000, "askingPrice": 14800000, "currency": "USD"},
                "expectedMathCheckStatus": "failed",
                "expectedRecommendation": "WALK AWAY",
                "expectedCrossDocumentConflicts": [
                    {
                        "metric": "revenue",
                        "period": "TTM",
                        "description": "Seller claimed revenue of $14,800,000 overstates bank-verified tax return revenue of $8,200,000 by $6,600,000 (44.6%).",
                        "severity": "critical"
                    }
                ]
            }
        },
        {
            "fileName": "TerraNova_Environmental_General_Ledger_Bank_Recon.xlsx",
            "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Balance Sheet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "FY2025", "rawValue": "$8,200,000"}
                ],
                "expectedRedFlags": ["$6.6M cash deficit vs billed invoices"],
                "expectedYellowFlags": [],
                "expectedRecommendation": "WALK AWAY"
            }
        },
        {
            "fileName": "TerraNova_Environmental_Tax_Form_1120_Audit.xlsx",
            "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Tax Return",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "FY2025", "rawValue": "$8,200,000"},
                    {"metric": "net_income", "normalizedValue": -410000, "period": "FY2025", "rawValue": "-$410,000"}
                ],
                "expectedRedFlags": ["-$410,000 operating loss on $8.2M revenue"],
                "expectedYellowFlags": [],
                "expectedRecommendation": "WALK AWAY"
            }
        },
        {
            "fileName": "TerraNova_Environmental_AR_Aging_Bad_Debt.xlsx",
            "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
            "fileType": "XLSX",
            "groundTruth": {
                "documentType": "Operating Metrics",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "TTM", "rawValue": "$8,200,000"}
                ],
                "expectedRedFlags": ["$3.34M in uncollectible AR write-offs (68% delinquent)"],
                "expectedYellowFlags": [],
                "expectedRecommendation": "WALK AWAY"
            }
        },
        {
            "fileName": "TerraNova_Environmental_Legal_Regulatory_Disclosures.docx",
            "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
            "fileType": "DOCX",
            "groundTruth": {
                "documentType": "Due Diligence Data Room Packet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [],
                "expectedRedFlags": ["$2.4M active EPA Superfund cleanup lien under federal consent decree"],
                "expectedYellowFlags": [],
                "expectedRecommendation": "WALK AWAY"
            }
        }
    ]
    for gt in gt_docs:
        gt_path = os.path.join(GT_DIR, f"packet3_terranova_{gt['fileName']}.json")
        with open(gt_path, "w") as f:
            json.dump(gt, f, indent=2)
            
    # Result Run file
    result_data = {
        "business": "TerraNova Environmental Hauling, LLC (Commercial Waste & Remediation)",
        "projectId": "project-terranova-environmental-walkaway",
        "evaluatedAt": "2026-08-19T20:00:00.000Z",
        "documents": [
            {
                "fileName": "TerraNova_Environmental_Seller_Teaser_CIM.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Due Diligence Data Room Packet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 14800000, "period": "TTM", "confidence": 1.0},
                    {"metric": "net_income", "normalizedValue": -410000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": [
                    "Severe revenue inflation: Seller claims $14.8M vs verified $8.2M bank cash ($6.6M deficit)",
                    "True operational loss of -$410,000 (insolvency risk)",
                    "Undisclosed $2.4M EPA Superfund cleanup lien",
                    "68% of AR >90 days delinquent ($3.3M bad debt write-off required)"
                ],
                "yellowFlags": [],
                "valuation": {"base_estimate": 0},
                "mathCheckStatus": "failed"
            },
            {
                "fileName": "TerraNova_Environmental_General_Ledger_Bank_Recon.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Balance Sheet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "FY2025", "confidence": 1.0}
                ],
                "redFlags": ["$6.6M cash deficit vs billed invoices"],
                "yellowFlags": [],
                "mathCheckStatus": "failed"
            },
            {
                "fileName": "TerraNova_Environmental_Tax_Form_1120_Audit.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Tax Return",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "FY2025", "confidence": 1.0},
                    {"metric": "net_income", "normalizedValue": -410000, "period": "FY2025", "confidence": 1.0}
                ],
                "redFlags": ["-$410,000 operating loss on $8.2M revenue"],
                "yellowFlags": [],
                "mathCheckStatus": "failed"
            },
            {
                "fileName": "TerraNova_Environmental_AR_Aging_Bad_Debt.xlsx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Operating Metrics",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [
                    {"metric": "revenue", "normalizedValue": 8200000, "period": "TTM", "confidence": 1.0}
                ],
                "redFlags": ["$3.34M in uncollectible AR write-offs (68% delinquent)"],
                "yellowFlags": [],
                "mathCheckStatus": "failed"
            },
            {
                "fileName": "TerraNova_Environmental_Legal_Regulatory_Disclosures.docx",
                "status": "completed",
                "modelUsed": "OpenAI 5.6 Terra",
                "detectedDocumentType": "Due Diligence Data Room Packet",
                "trafficLight": "RED",
                "riskLevel": "CRITICAL",
                "financialFacts": [],
                "redFlags": ["$2.4M active EPA Superfund cleanup lien under federal consent decree"],
                "yellowFlags": [],
                "mathCheckStatus": "failed"
            }
        ]
    }
    with open(os.path.join(RESULTS_DIR, "packet3_terranova_environmental_actual_run.json"), "w") as f:
        json.dump(result_data, f, indent=2)

if __name__ == "__main__":
    print("Generating Packet 1: Vanguard Medical Logistics (PROCEED)...")
    build_packet_1()
    print("Generating Packet 2: Apex Precision Dynamics (RENEGOTIATE)...")
    build_packet_2()
    print("Generating Packet 3: TerraNova Environmental (WALK AWAY)...")
    build_packet_3()
    print("All 3 deal packets (15 documents + ground truth + run results) generated successfully!")
